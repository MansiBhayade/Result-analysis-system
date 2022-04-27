
import pandas as pd
df = pd.read_excel('result.xlsx')
df.columns = df.iloc[0]
df = df.iloc[1:].reset_index(drop=True)
df1=df.drop(index=df.index[0])

std_no=df1.shape[0]


#pass students
cgpi = df1.get("GPA")
prows = df1.shape[0]
pass_std=prows-1
p_pass_std=(pass_std*100)/std_no
pass_s=str(pass_std)
fail=1
p_f=(fail*100)/std_no
c_1=92
c_2=20
dist=16
p_c_1=(c_1*100)/std_no
p_c_2=(c_2*100)/std_no
p_dist=(dist*100)/std_no




s=str(std_no)           #total student
pass_s=str(pass_std)    #passed stuent
p_pass_s=str(p_pass_std)  #per passed student
dist_std=str(dist)
c_1_std=str(c_1)
c_2_std=str(c_2)
p_c_1_std=str(p_c_1)
p_c_2_std=str(p_c_2)
dist_std=str(dist)
p_dist_std=str(p_dist)
f_std=str(fail)
p_f_std=str(p_f)


from xml.dom.minidom import Document
import docx

doc=docx.Document()

section= doc.sections[0]
header = section.header

header_para= header.paragraphs[0]
header_para.text="DATTA MEGHE COLLEGE OF ENGINEERING"
doc.add_heading('Result analysis')

footer= section.footer
footer_para =footer.paragraphs[0]
footer_para.text="Made by SE A-10,11,17"

data=[
    ['Appeared',s,s,s,s,s,s,s,s,s],
    ['Passed', pass_s, pass_s,pass_s, pass_s ,pass_s,pass_s, pass_s, pass_s,pass_s,pass_s],
    ['Pass %',p_pass_s   ,p_pass_s   ,p_pass_s   ,p_pass_s   ,p_pass_s   ,p_pass_s   ,p_pass_s   ,p_pass_s   ,p_pass_s   ,p_pass_s   ],
    ['Distinction',dist_std  ,dist_std  ,dist_std  ,dist_std  ,dist_std  ,dist_std  ,dist_std  ,dist_std  ,dist_std  ],
    ['1st class',c_1_std ,c_1_std ,c_1_std ,c_1_std ,c_1_std ,c_1_std ,c_1_std ,c_1_std ,c_1_std ],
    ['1st class %',p_c_1_std ,p_c_1_std ,p_c_1_std,p_c_1_std ,p_c_1_std ,p_c_1_std ,p_c_1_std ,p_c_1_std ,p_c_1_std ],
    ['II class(d]',c_2_std ,c_2_std ,c_2_std ,c_2_std ,c_2_std ,c_2_std ,c_2_std ,c_2_std ,c_2_std ],
    ['II class %',p_c_2_std ,p_c_2_std ,p_c_2_std ,p_c_2_std ,p_c_2_std ,p_c_2_std ,p_c_2_std ,p_c_2_std ,p_c_2_std ],
    ['Absent',f_std ,f_std ,f_std ,f_std , f_std,f_std ,f_std ,f_std ,f_std ],
    ['Fail',f_std , f_std, f_std,f_std ,f_std ,f_std ,f_std ,f_std ,f_std ],
    ['Fail %',p_f_std ,p_f_std ,p_f_std ,p_f_std ,p_f_std ,p_f_std ,p_f_std ,p_f_std ,p_f_std ]
]

table =doc.add_table(rows=1,cols=10)
table.style='Table Grid'
row =table.rows[0].cells
row[0].text ='Category/Subject'
row[1].text ='DS&IP'
row[2].text ='MCC'
row[3].text ='AI'
row[4].text ='DLC-II'
row[5].text ='ILOC-1'
row[6].text ='MADL'
row[7].text ='CL-1'
row[8].text ='MJ-I'
row[9].text ='TOTAL'

for item in data:
    row=table.add_row().cells
    row[0].text =item[0]
    row[1].text =item[1]
    row[2].text =item[2]
    row[3].text =item[3]
    row[4].text =item[4]
    row[5].text =item[5]
    row[6].text =item[6]
    row[7].text =item[7]
    row[8].text =item[8]
    row[9].text =item[9]

doc.save('result-analysis.docx')

